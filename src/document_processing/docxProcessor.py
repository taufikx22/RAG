import docx
import logging
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List

from src.document_processing.base import DocumentProcessor

logger = logging.getLogger(__name__)

class DocxProcessor(DocumentProcessor):
    def process(self, file_path: Path) -> Dict[str, Any]:
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() != '.docx':
            logger.error(f"Invalid file type: {file_path.suffix}. Expected .docx")
            raise ValueError(f"Invalid file type: {file_path.suffix}. Expected .docx")
        
        try:
            doc= docx.Document(file_path)
            content= ""
            for para in doc.paragraphs:
                content += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + "\n"

            metadata = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'last_modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'hash': self._calculate_hash(file_path),
            }

            try:
                doc_id = self._generate_doc_id(file_path)
                metadata.update({
                    'author': doc.core_properties.author,
                    'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    'title': doc.core_properties.title,
                    'subject': doc.core_properties.subject,
                    'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                })
            except Exception as e:
                logger.warning(f"Error extracting metadata from {file_path}: {e}")
                doc_id = None

            doc_id= hashlib.md5(file_path.name.encode()).hexdigest()
            metadata['doc_id'] = doc_id

            return {
                'content': content.strip(),
                'metadata': metadata,
                'doc_id': doc_id
            }
        
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise e